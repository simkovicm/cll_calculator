"""
CLL Response Assessment Calculator v2
iwCLL 2018 Guidelines — s historií hodnocení, grafy a exportem
"""

import streamlit as st
from datetime import date
import json
import io

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="CLL Kalkulátor",
    page_icon="🩺",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

/* Sidebar styling */
section[data-testid="stSidebar"] { background: #0f172a; }
section[data-testid="stSidebar"] * { color: #e2e8f0 !important; }
section[data-testid="stSidebar"] .stSelectbox label,
section[data-testid="stSidebar"] .stNumberInput label,
section[data-testid="stSidebar"] .stTextInput label { color: #94a3b8 !important; font-size: 0.78rem !important; }

.main-header {
    background: linear-gradient(135deg, #1e3a5f 0%, #102a4c 100%);
    padding: 1.5rem 2rem;
    border-radius: 12px;
    margin-bottom: 1.5rem;
    border-left: 5px solid #3b82f6;
}
.main-header h1 { color: white; font-size: 1.6rem; font-weight: 700; margin: 0; }
.main-header p  { color: #94a3b8; margin: 0.3rem 0 0; font-size: 0.85rem; }

/* Patient card */
.patient-card {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 1rem 1.2rem;
    margin-bottom: 1rem;
    font-size: 0.85rem;
}
.patient-card b { color: #1e3a5f; }

/* Section headers */
.sec-a {
    background: #eff6ff; border-left: 4px solid #3b82f6;
    padding: 7px 14px; border-radius: 0 8px 8px 0;
    font-weight: 600; color: #1e40af; font-size: 0.88rem; margin-bottom: 0.6rem;
}
.sec-b {
    background: #f0fdf4; border-left: 4px solid #22c55e;
    padding: 7px 14px; border-radius: 0 8px 8px 0;
    font-weight: 600; color: #15803d; font-size: 0.88rem; margin-bottom: 0.6rem;
}
.sec-warn {
    background: #fff7ed; border-left: 4px solid #f59e0b;
    padding: 7px 14px; border-radius: 0 8px 8px 0;
    font-weight: 600; color: #92400e; font-size: 0.88rem; margin-bottom: 0.6rem;
}

/* Result box */
.result-box {
    border-radius: 14px; padding: 1.8rem; text-align: center;
    margin: 0.5rem 0 1rem; border: 2px solid;
}
.result-label { font-size: 0.7rem; text-transform: uppercase; letter-spacing: 3px; opacity: 0.7; margin-bottom: 0.4rem; }
.result-value { font-size: 3.5rem; font-weight: 800; line-height: 1; margin: 0.3rem 0; }
.result-desc  { font-size: 0.82rem; line-height: 1.5; margin-top: 0.5rem; }
.result-CR  { background:#dcfce7; border-color:#16a34a; color:#14532d; }
.result-CRi { background:#dbeafe; border-color:#2563eb; color:#1e3a8a; }
.result-PR  { background:#dbeafe; border-color:#3b82f6; color:#1e40af; }
.result-PRL { background:#f3e8ff; border-color:#9333ea; color:#581c87; }
.result-nPR { background:#e0f2fe; border-color:#0284c7; color:#0c4a6e; }
.result-SD  { background:#fefce8; border-color:#ca8a04; color:#713f12; }
.result-PD  { background:#fee2e2; border-color:#dc2626; color:#7f1d1d; }

/* Breakdown */
.bd-row {
    display:flex; justify-content:space-between; align-items:center;
    padding:7px 4px; border-bottom:1px solid #f1f5f9; font-size:0.82rem;
}
.bd-name { color:#475569; }
.bd-val  { font-weight:500; color:#1e293b; }
.bmet    { background:#dcfce7; color:#166534; padding:2px 9px; border-radius:4px; font-size:0.72rem; font-weight:600; }
.bnotmet { background:#fee2e2; color:#991b1b; padding:2px 9px; border-radius:4px; font-size:0.72rem; font-weight:600; }
.bwarn   { background:#fef9c3; color:#713f12; padding:2px 9px; border-radius:4px; font-size:0.72rem; font-weight:600; }
.bna     { background:#f1f5f9; color:#64748b; padding:2px 9px; border-radius:4px; font-size:0.72rem; }
.note-w  { background:#fef9c3; border:1px solid #fcd34d; border-radius:8px; padding:0.8rem; font-size:0.8rem; color:#713f12; line-height:1.6; margin-top:0.5rem; }
.info-bar { background:#eff6ff; border:1px solid #bfdbfe; border-radius:7px; padding:7px 12px; font-size:0.8rem; color:#1e40af; margin-bottom:0.8rem; }

/* History table */
.hist-table { width:100%; border-collapse:collapse; font-size:0.82rem; }
.hist-table th { background:#1e3a5f; color:white; padding:7px 10px; text-align:left; }
.hist-table td { padding:6px 10px; border-bottom:1px solid #e2e8f0; }
.hist-table tr:nth-child(even) td { background:#f8fafc; }

/* Chip */
.chip-CR  { background:#dcfce7; color:#166534; border-radius:20px; padding:2px 10px; font-weight:700; font-size:0.8rem; }
.chip-PR  { background:#dbeafe; color:#1e40af; border-radius:20px; padding:2px 10px; font-weight:700; font-size:0.8rem; }
.chip-PRL { background:#f3e8ff; color:#6b21a8; border-radius:20px; padding:2px 10px; font-weight:700; font-size:0.8rem; }
.chip-CRi { background:#dbeafe; color:#1e3a8a; border-radius:20px; padding:2px 10px; font-weight:700; font-size:0.8rem; }
.chip-nPR { background:#e0f2fe; color:#0c4a6e; border-radius:20px; padding:2px 10px; font-weight:700; font-size:0.8rem; }
.chip-SD  { background:#fefce8; color:#713f12; border-radius:20px; padding:2px 10px; font-weight:700; font-size:0.8rem; }
.chip-PD  { background:#fee2e2; color:#7f1d1d; border-radius:20px; padding:2px 10px; font-weight:700; font-size:0.8rem; }
</style>
""", unsafe_allow_html=True)

# ── Session state ─────────────────────────────────────────────────────────────
if "history" not in st.session_state:
    st.session_state.history = []   # list of dicts
if "result" not in st.session_state:
    st.session_state.result = None

# ══════════════════════════════════════════════════════════════════════════════
#  CORE LOGIC
# ══════════════════════════════════════════════════════════════════════════════
def compute(p):
    """All iwCLL 2018 logic. p = dict of params. Returns result dict."""

    def v(k): return p.get(k)

    spd_pct = None
    baseline_spd = p.get("baseline_spd", 0) or 0
    current_spd  = p.get("current_spd", 0) or 0
    has_any_ln   = p.get("has_any_ln", False)
    all_nodes_small = p.get("all_nodes_small", True)

    if has_any_ln and baseline_spd > 0:
        spd_pct = (current_spd - baseline_spd) / baseline_spd * 100

    # ── PD ────────────────────────────────────────────────────────────────────
    is_pd, pd_r = False, []
    if spd_pct is not None and spd_pct >= 50: is_pd=True; pd_r.append("SPD vzrostlo ≥50%")
    if v("new_lesions"):  is_pd=True; pd_r.append("Nové lymfatické léze")
    if v("new_cytopenia"): is_pd=True; pd_r.append("Nová cytopenie")
    sp_b, sp_c = v("spleen_base"), v("spleen_curr")
    if sp_b and sp_c and sp_c > sp_b*1.5: is_pd=True; pd_r.append("Slezina vzrostla ≥50%")
    hb, hc = v("hgb_base"), v("hgb_curr")
    if hb and hc and (hb-hc)>=2: is_pd=True; pd_r.append("Hb poklesl ≥2 g/dL")
    pb, pc = v("plt_base"), v("plt_curr")
    if pb and pc and pc < pb*0.5: is_pd=True; pd_r.append("Trombocyty poklesly ≥50%")
    if v("marrow_key")=="increasing": is_pd=True; pd_r.append("Nárůst CLL buněk v dřeni")

    # ── Group A ───────────────────────────────────────────────────────────────
    A_met=0; A_abn=0; a_items=[]; lymphocytosis=False

    # LN
    ln_abn = has_any_ln and baseline_spd>0
    if has_any_ln:
        if ln_abn and spd_pct is not None:
            A_abn+=1
            s = "met" if spd_pct<=-50 else "notmet"
            if s=="met": A_met+=1
            a_items.append(("Lymf. uzliny (SPD)", f"{spd_pct:+.1f}% (BL: {baseline_spd:.2f} cm²)", s, None))
        else:
            a_items.append(("Lymf. uzliny (SPD)", f"SPD: {current_spd:.2f} cm²", "na", "Bez baseline"))
    else:
        a_items.append(("Lymf. uzliny (SPD)", "—", "na", None))

    # Lymphocytes
    lb, lc = v("lymph_base"), v("lymph_curr")
    lb_abn = lb is not None and lb>5
    if lc is not None:
        if lb_abn:
            A_abn+=1
            lp = (lb-lc)/lb*100
            if lc<4: A_met+=1; ls="met"
            elif lp>=50: A_met+=1; ls="met"
            elif lc>=5: lymphocytosis=True; ls="warn"
            else: ls="notmet"
            a_items.append(("Lymfocyty", f"{lc} ×10⁹/L", ls, None))
        elif lb is not None:
            a_items.append(("Lymfocyty", f"{lc} ×10⁹/L", "na", "Normální při baseline"))
        else:
            ls2 = "warn" if lc>=5 else "na"
            if lc>=5: lymphocytosis=True
            a_items.append(("Lymfocyty", f"{lc} ×10⁹/L", ls2, "Baseline nezadán"))
    else:
        a_items.append(("Lymfocyty", "—", "na", None))

    # Spleen
    sp_abn = sp_b is not None and sp_b>=13
    if sp_c is not None:
        if sp_abn:
            A_abn+=1
            spp = (sp_b-sp_c)/sp_b*100
            ss = "met" if (spp>=50 or sp_c<13) else "notmet"
            if ss=="met": A_met+=1
            a_items.append(("Slezina", f"{sp_c} cm", ss, None))
        else:
            a_items.append(("Slezina", f"{sp_c} cm", "na", "Normální při baseline (<13 cm)"))
    else:
        a_items.append(("Slezina", "—", "na", None))

    # Liver
    lv_abn = v("liver_base_abn")
    lv_curr_n = v("liver_curr_normal")
    if lv_abn:
        A_abn+=1
        lvs = "met" if lv_curr_n else "notmet"
        if lvs=="met": A_met+=1
    else:
        lvs = "na"
    a_items.append(("Játra", "Normální" if lv_curr_n else "Hepatomegalie", lvs,
                    "Normální při baseline" if not lv_abn else None))

    # Constitutional
    cs_abn = v("const_base_abn")
    cs_n   = v("const_curr_none")
    if cs_abn:
        A_abn+=1
        css = "met" if cs_n else "notmet"
        if css=="met": A_met+=1
    else:
        css = "na" if cs_n else "notmet"
    a_items.append(("Konstituční příznaky", "Žádné" if cs_n else "Přítomny", css,
                    "Nepřítomny při baseline" if not cs_abn else None))

    # ── Group B ───────────────────────────────────────────────────────────────
    B_met=0; B_abn=0; b_items=[]

    # Platelets
    pb_abn = pb is not None and pb<100
    if pc is not None:
        if pb_abn:
            B_abn+=1
            pts = "met" if (pc>=100 or (pc-pb)/pb*100>=50) else "notmet"
            if pts=="met": B_met+=1
        else:
            pts = "na" if (pb is None or pc>=100) else "notmet"
        b_items.append(("Trombocyty", f"{pc} ×10⁹/L", pts,
                        "Normální při baseline" if (pb is not None and not pb_abn) else None))
    else:
        b_items.append(("Trombocyty", "—", "na", None))

    # Hgb
    hb_abn = hb is not None and hb<11
    if hc is not None:
        if hb_abn:
            B_abn+=1
            hbs = "met" if (hc>=11 or (hc-hb)/hb*100>=50) else "notmet"
            if hbs=="met": B_met+=1
        else:
            hbs = "na" if (hb is None or hc>=11) else "notmet"
        b_items.append(("Hemoglobin", f"{hc} g/dL", hbs,
                        "Normální při baseline" if (hb is not None and not hb_abn) else None))
    else:
        b_items.append(("Hemoglobin", "—", "na", None))

    # ANC
    ab, ac = v("anc_base"), v("anc_curr")
    ab_abn = ab is not None and ab<1500
    if ac is not None:
        if ab_abn:
            B_abn+=1
            acs = "met" if (ac>=1500 or (ac-ab)/ab*100>=50) else "notmet"
            if acs=="met": B_met+=1
        else:
            acs = "na" if (ab is None or ac>=1500) else "notmet"
        b_items.append(("ANC (neutrofily)", f"{ac} ×10⁹/L", acs,
                        "Normální při baseline" if (ab is not None and not ab_abn) else None))
    else:
        b_items.append(("ANC (neutrofily)", "—", "na", None))

    # Marrow
    mk = v("marrow_key") or "not_done"
    if mk=="normal":      B_met+=1; B_abn+=1; ms="met";    mn=None
    elif mk=="not_done":  B_met+=1; B_abn+=1; ms="warn";   mn="Nebyla provedena — počítá se jako splněno (iwCLL 2018)"
    elif mk in ("cll_cells","nodules"): B_abn+=1; ms="warn"; mn=None
    elif mk=="increasing": B_abn+=1; ms="notmet"; mn=None
    else: ms="na"; mn=None
    b_items.append(("Kostní dřeň", {
        "not_done":"Nebyla provedena","normal":"Normální","cll_cells":"CLL buňky/noduly",
        "nodules":"Lymfoidní noduly","increasing":"Nárůst CLL"
    }.get(mk,"—"), ms, mn))

    B_all_norm = B_abn==0
    total_abn  = A_abn+B_abn
    single     = total_abn<=1
    pA_thr     = 1 if single else 2
    pB_thr     = 0 if single else 1
    B_sat      = B_met>=1 or B_all_norm
    pr_met     = (A_met>=1 or B_met>=1) if single else (A_met>=pA_thr and B_sat)

    # ── CR criteria ───────────────────────────────────────────────────────────
    cr_nodes  = not has_any_ln or all_nodes_small
    cr_spleen = sp_c is None or sp_c<13
    cr_liver  = lv_curr_n
    cr_const  = cs_n
    cr_lymph  = lc is not None and lc<4
    cr_plt    = pc is not None and pc>=100
    cr_hgb    = hc is not None and hc>=11
    cr_anc    = ac is not None and ac>=1500
    cr_A = cr_nodes and cr_spleen and cr_liver and cr_const and cr_lymph
    cr_B = cr_plt and cr_hgb and cr_anc
    cr_marrow_ok   = mk=="normal" and (v("marrow_pct") is None or v("marrow_pct")<30)
    cr_marrow_done = mk!="not_done"

    # ── Determine response ────────────────────────────────────────────────────
    notes=[]; response="SD"; desc=""
    if is_pd:
        response="PD"; desc="Progresivní onemocnění — splněno ≥1 kritérium pro progresi"
        for r in pd_r: notes.append("Kritérium: "+r)
    elif cr_A and cr_B and cr_marrow_ok and cr_marrow_done:
        response="CR"; desc="Kompletní remise — všechna kritéria A i B splněna, KD normální"
    elif cr_A and cr_B and not cr_marrow_done:
        response="CR"; desc="Kompletní remise — splněna A i B kritéria (nutno potvrdit biopsií KD)"
        notes.append("Pro potvrzení CR je nutná biopsie KD: ≤30% lymfocytů, bez lymfoidních nodulů")
    elif cr_A and cr_B and mk=="nodules":
        response="nPR"; desc="Nodulární parciální remise — CR kritéria splněna v krvi, lymfoidní noduly v KD"
    elif cr_A and not (cr_plt and cr_hgb and cr_anc):
        response="CRi"; desc="CR s neúplnou obnovou dřeně — splněna A kritéria, přetrvává cytopenie"
    elif pr_met and not lymphocytosis:
        response="PR"
        if single and total_abn==1:
            desc="Parciální remise — pravidlo jediného parametru (1 abnormální → stačí zlepšení 1 parametru)"
        else:
            desc=f"Parciální remise — zlepšení ≥{pA_thr} param. skupiny A a ≥1 param. skupiny B"
    elif pr_met and lymphocytosis:
        response="PR-L"; desc=f"Parciální remise s lymfocytózou — splněna kritéria PR, lymfocyty zvýšeny redistribucí"
        notes.append("Lymfocytóza u inhibitorů kináz (ibrutinib aj.) sama o sobě není PD — redistribuce (iwCLL 2018 sec. 5.3.4)")
    else:
        response="SD"; desc=f"Stabilní onemocnění — nesplněna kritéria CR/PR, bez progrese (A: {A_met}/{A_abn}, B: {B_met}/{B_abn})"
        if A_abn==0 and B_abn==0:
            notes.append("Žádné parametry nebyly patologické při baseline — zkontrolujte zadané hodnoty")

    # ── Doc text ──────────────────────────────────────────────────────────────
    resp_labels = {
        "CR":"kompletní remise (CR)","CRi":"CR s neúplnou obnovou dřeně (CRi)",
        "PR":"parciální remise (PR)","PR-L":"parciální remise s lymfocytózou (PR-L)",
        "nPR":"nodulární parciální remise (nPR)","PD":"progrese onemocnění (PD)",
        "SD":"stabilní onemocnění (SD)"
    }
    today_s = date.today().strftime("%-d. %-m. %Y")
    doc = [f"Hodnocení odpovědi ({today_s}): Jedná se o {resp_labels.get(response,response)}."]
    reas=[]
    if spd_pct is not None and has_any_ln:
        if spd_pct<=-50: reas.append(f"pokles SPD o {abs(spd_pct):.0f}% ({baseline_spd:.2f}→{current_spd:.2f} cm²)")
        elif spd_pct>=50: reas.append(f"nárůst SPD o {spd_pct:.0f}% (progrese)")
        else: reas.append(f"změna SPD o {spd_pct:+.0f}%")
    if lc and lb_abn:
        if lc<4: reas.append(f"normalizace lymfocytů ({lb}→{lc} ×10⁹/L)")
        elif lymphocytosis: reas.append(f"redistribuční lymfocytóza ({lc} ×10⁹/L)")
    if sp_c and sp_abn:
        spp2=(sp_b-sp_c)/sp_b*100
        if spp2>=50 or sp_c<13: reas.append(f"redukce splenomegalie ({sp_b}→{sp_c} cm)")
    if lv_abn and lv_curr_n: reas.append("normalizace hepatomegalie")
    if cs_abn and cs_n: reas.append("vymizení konstitučních příznaků")
    if pc and pb_abn and pc>=100: reas.append(f"normalizace trombocytů ({pb}→{pc} ×10⁹/L)")
    if hc and hb_abn and hc>=11:  reas.append(f"normalizace hemoglobinu ({hb}→{hc} g/dL)")
    if ac and ab_abn and ac>=1500: reas.append(f"normalizace neutrofilů ({ab}→{ac} ×10⁹/L)")
    mk_doc={"normal":"normální KD","cll_cells":"CLL buňky v KD","nodules":"lymfoidní noduly v KD","not_done":"KD nebyla provedena"}
    if mk in mk_doc: reas.append(mk_doc[mk])
    for r in pd_r: reas.append(r)
    if reas: doc.append("Odůvodnění: "+", ".join(reas)+".")
    doc.append("Hodnocení dle iwCLL 2018 guidelines.")

    return {
        "response": response, "desc": desc, "notes": notes,
        "a_items": a_items, "b_items": b_items,
        "A_met": A_met, "A_abn": A_abn, "B_met": B_met, "B_abn": B_abn,
        "pA_thr": pA_thr, "pB_thr": pB_thr,
        "B_all_norm": B_all_norm, "single": single,
        "doc_text": " ".join(doc),
        "spd_pct": spd_pct,
        "baseline_spd": baseline_spd, "current_spd": current_spd,
    }


# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR — patient info + inputs
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("### 🩺 CLL Kalkulátor")
    st.markdown("---")

    # Patient info
    st.markdown("**Pacient / Studie**")
    pat_id    = st.text_input("ID / jméno pacienta", placeholder="např. CLL-001", key="pat_id")
    study_name= st.text_input("Název studie / protokolu", placeholder="např. MAJIC", key="study")
    tp_label  = st.text_input("Časový bod", placeholder="např. C6 screening, M12", key="tp_label")
    eval_date = st.date_input("Datum hodnocení", value=date.today(), key="eval_date")

    st.markdown("---")

    # ── LN mode ───────────────────────────────────────────────────────────────
    st.markdown("**🔵 Lymfatické uzliny — SPD**")
    ln_mode = st.radio("Způsob zadání", ["Jednotlivé uzliny", "Přímé SPD"], key="ln_mode", horizontal=True)

    baseline_spd=0.0; current_spd=0.0; has_any_ln=False; all_nodes_small=True

    if ln_mode=="Jednotlivé uzliny":
        n_nodes=st.number_input("Počet uzlin",1,6,1,key="n_nodes")
        for i in range(int(n_nodes)):
            with st.expander(f"Uzlina {i+1}", expanded=(i==0)):
                c1,c2=st.columns(2)
                b1=c1.number_input("BL d₁",0.0,step=0.1,key=f"b1_{i}",label_visibility="visible")
                b2=c2.number_input("BL d₂",0.0,step=0.1,key=f"b2_{i}",label_visibility="visible")
                c3,c4=st.columns(2)
                c1v=c3.number_input("Akt. d₁",0.0,step=0.1,key=f"c1_{i}",label_visibility="visible")
                c2v=c4.number_input("Akt. d₂",0.0,step=0.1,key=f"c2_{i}",label_visibility="visible")
                ppdb=(b1*b2) if b1>0 and b2>0 else None
                ppdc=(c1v*c2v) if c1v>0 and c2v>0 else None
                if ppdb: baseline_spd+=ppdb; has_any_ln=True
                if ppdc: current_spd+=ppdc
                if c1v>1.5 or c2v>1.5: all_nodes_small=False
    else:
        b_spd=st.number_input("SPD baseline (cm²)",0.0,step=0.01,key="spd_base_d")
        c_spd=st.number_input("SPD aktuálně (cm²)",0.0,step=0.01,key="spd_curr_d")
        if b_spd>0:
            has_any_ln=True; baseline_spd=b_spd; current_spd=c_spd
            all_nodes_small=(c_spd==0)

    if has_any_ln and baseline_spd>0:
        pct=(current_spd-baseline_spd)/baseline_spd*100
        col="🟢" if pct<=-50 else ("🔴" if pct>=50 else "🟡")
        st.caption(f"{col} SPD: {baseline_spd:.2f} → {current_spd:.2f} cm² ({pct:+.1f}%)")

    st.markdown("---")
    st.markdown("**🔵 Slezina a játra**")
    c1,c2=st.columns(2)
    spleen_base=c1.number_input("Slezina BL (cm)",0.0,step=0.1,key="sp_b")
    spleen_curr=c2.number_input("Slezina akt. (cm)",0.0,step=0.1,key="sp_c")
    liver_base_abn  = st.selectbox("Játra — baseline", ["Normální","Hepatomegalie"],key="lv_b")=="Hepatomegalie"
    liver_curr_normal= st.selectbox("Játra — aktuálně",["Normální","Hepatomegalie"],key="lv_c")=="Normální"

    st.markdown("**🔵 Konstituční příznaky**")
    c1,c2=st.columns(2)
    const_base_abn = c1.selectbox("Baseline",["Žádné","Přítomny"],key="cs_b")=="Přítomny"
    const_curr_none= c2.selectbox("Aktuálně",["Žádné","Přítomny"],key="cs_c")=="Žádné"

    st.markdown("**🔵 Lymfocyty (×10⁹/L)**")
    c1,c2=st.columns(2)
    lymph_base=c1.number_input("Baseline",0.0,step=0.1,key="lb")
    lymph_curr=c2.number_input("Aktuálně",0.0,step=0.1,key="lc")

    st.markdown("---")
    st.markdown("**🟢 Krevní obraz**")
    c1,c2=st.columns(2)
    plt_base=c1.number_input("PLT BL (×10⁹/L)",0.0,step=1.0,key="pb")
    plt_curr=c2.number_input("PLT akt.",0.0,step=1.0,key="pc")
    c1,c2=st.columns(2)
    hgb_base=c1.number_input("Hgb BL (g/dL)",0.0,step=0.1,key="hb")
    hgb_curr=c2.number_input("Hgb akt.",0.0,step=0.1,key="hc")
    c1,c2=st.columns(2)
    anc_base=c1.number_input("ANC BL (×10⁹/L)",0.0,step=1.0,key="ab")
    anc_curr=c2.number_input("ANC akt.",0.0,step=1.0,key="ac")

    st.markdown("**🟢 Kostní dřeň**")
    marrow_opts={
        "Nebyla provedena":"not_done",
        "Normální":"normal",
        "CLL buňky/B-lymf. noduly":"cll_cells",
        "Lymfoidní noduly":"nodules",
        "Nárůst CLL buněk":"increasing"
    }
    marrow_sel=st.selectbox("Biopsie KD",list(marrow_opts.keys()),key="marrow")
    marrow_key=marrow_opts[marrow_sel]
    marrow_pct_v=st.number_input("% lymfocytů v KD (pro CR)",0,100,step=1,key="mpct")

    st.markdown("---")
    st.markdown("**⚠️ Nové léze**")
    c1,c2=st.columns(2)
    new_lesions = c1.selectbox("Nové uzliny/léze",["Ne","Ano"],key="nl")=="Ano"
    new_cytopenia=c2.selectbox("Nová cytopenie",["Ne","Ano"],key="nc")=="Ano"

    st.markdown("---")
    calc_btn=st.button("🔬 Vyhodnotit", type="primary", use_container_width=True)


# ══════════════════════════════════════════════════════════════════════════════
#  HELPER — build params dict
# ══════════════════════════════════════════════════════════════════════════════
def _val(v): return v if v and v>0 else None

def build_params():
    return {
        "baseline_spd": baseline_spd, "current_spd": current_spd,
        "has_any_ln": has_any_ln, "all_nodes_small": all_nodes_small,
        "spleen_base": _val(spleen_base), "spleen_curr": _val(spleen_curr),
        "liver_base_abn": liver_base_abn, "liver_curr_normal": liver_curr_normal,
        "const_base_abn": const_base_abn, "const_curr_none": const_curr_none,
        "lymph_base": _val(lymph_base), "lymph_curr": _val(lymph_curr),
        "plt_base": _val(plt_base), "plt_curr": _val(plt_curr),
        "hgb_base": _val(hgb_base), "hgb_curr": _val(hgb_curr),
        "anc_base": _val(anc_base), "anc_curr": _val(anc_curr),
        "marrow_key": marrow_key, "marrow_pct": _val(float(marrow_pct_v)) if marrow_pct_v else None,
        "new_lesions": new_lesions, "new_cytopenia": new_cytopenia,
    }


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN PANEL
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="main-header">
  <h1>🩺 CLL Response Assessment Calculator</h1>
  <p>Hodnocení odpovědi na léčbu dle iwCLL 2018 Guidelines &nbsp;|&nbsp; Parametry zadávejte v levém panelu</p>
</div>
""", unsafe_allow_html=True)

# ── Tabs ──────────────────────────────────────────────────────────────────────
tab_result, tab_history, tab_export = st.tabs(["📊 Aktuální hodnocení", "📅 Historie hodnocení", "📄 Export"])

# ─────────────────────────────────────────────────────────────────────────────
#  TAB 1 — Result
# ─────────────────────────────────────────────────────────────────────────────
with tab_result:
    # Patient card
    if pat_id or study_name or tp_label:
        parts=[]
        if pat_id: parts.append(f"<b>Pacient:</b> {pat_id}")
        if study_name: parts.append(f"<b>Studie:</b> {study_name}")
        if tp_label: parts.append(f"<b>Timepoint:</b> {tp_label}")
        parts.append(f"<b>Datum:</b> {eval_date.strftime('%-d. %-m. %Y')}")
        st.markdown(f'<div class="patient-card">{" &nbsp;|&nbsp; ".join(parts)}</div>', unsafe_allow_html=True)

    # Live calculation — always compute on any input change
    params = build_params()
    res = compute(params)
    st.session_state.result = res

    # Result card
    css_k = {"CR":"CR","CRi":"CRi","PR":"PR","PR-L":"PRL","nPR":"nPR","PD":"PD","SD":"SD"}.get(res["response"],"SD")
    st.markdown(f"""
    <div class="result-box result-{css_k}">
      <div class="result-label">Hodnocení odpovědi dle iwCLL 2018</div>
      <div class="result-value">{res["response"]}</div>
      <div class="result-desc">{res["desc"]}</div>
    </div>
    """, unsafe_allow_html=True)

    if calc_btn:
        # Save to history
        entry = {
            "date": eval_date.strftime("%-d. %-m. %Y"),
            "tp": tp_label or "—",
            "response": res["response"],
            "pat_id": pat_id or "—",
            "study": study_name or "—",
            "spd_pct": res["spd_pct"],
            "lymph_curr": _val(lymph_curr),
            "plt_curr": _val(plt_curr),
            "hgb_curr": _val(hgb_curr),
            "doc_text": res["doc_text"],
        }
        st.session_state.history.append(entry)
        st.success("✅ Hodnocení uloženo do historie")

    # Breakdown
    with st.expander("📋 Detailní rozbor parametrů", expanded=True):
        st.markdown(f"""
        <div class="info-bar">
          Patologické při baseline: <b>A: {res["A_abn"]} | B: {res["B_abn"]}</b>
          {"&nbsp;&nbsp;<span style='color:#92400e'>⚠ Pravidlo jediného parametru</span>" if res["single"] else ""}
          &nbsp;&nbsp;|&nbsp;&nbsp; Splněná zlepšení: <b>A: {res["A_met"]}/{res["A_abn"]} | B: {res["B_met"]}/{res["B_abn"]}</b>
        </div>
        """, unsafe_allow_html=True)

        bl_map = {"met":"bmet","notmet":"bnotmet","warn":"bwarn","na":"bna"}
        bl_lab = {"met":"✓ Zlepšeno","notmet":"✗ Nezlepšeno","warn":"⚠ Pozor","na":"— N/A"}

        st.markdown("**Skupina A**")
        for name,val,status,note in res["a_items"]:
            note_h = f'<br><small style="color:#92400e">{note}</small>' if note else ""
            st.markdown(f"""<div class="bd-row">
              <span class="bd-name">{name}{note_h}</span>
              <span><span class="bd-val">{val}</span>&nbsp;&nbsp;<span class="{bl_map[status]}">{bl_lab[status]}</span></span>
            </div>""", unsafe_allow_html=True)
        st.markdown(f"<div style='font-size:0.78rem;color:#1e40af;padding:5px 0;font-weight:600'>A — patologické: {res['A_abn']} | zlepšeno: {res['A_met']} | požadováno: {res['pA_thr']}</div>", unsafe_allow_html=True)

        st.markdown("**Skupina B**")
        for name,val,status,note in res["b_items"]:
            note_h = f'<br><small style="color:#92400e">{note}</small>' if note else ""
            st.markdown(f"""<div class="bd-row">
              <span class="bd-name">{name}{note_h}</span>
              <span><span class="bd-val">{val}</span>&nbsp;&nbsp;<span class="{bl_map[status]}">{bl_lab[status]}</span></span>
            </div>""", unsafe_allow_html=True)
        auto=""
        if res["B_all_norm"]: auto=" <span style='color:#64748b'>(vše normální → automaticky splněno)</span>"
        st.markdown(f"<div style='font-size:0.78rem;color:#15803d;padding:5px 0;font-weight:600'>B — patologické: {res['B_abn']} | zlepšeno: {res['B_met']} | požadováno: {res['pB_thr']}{auto}</div>", unsafe_allow_html=True)

    if res["notes"]:
        st.markdown('<div class="note-w">'+"<br><br>".join(["⚠ "+n for n in res["notes"]])+'</div>', unsafe_allow_html=True)

    # Doc text
    st.markdown("**📝 Text pro dokumentaci**")
    st.code(res["doc_text"], language=None)

# ─────────────────────────────────────────────────────────────────────────────
#  TAB 2 — History
# ─────────────────────────────────────────────────────────────────────────────
with tab_history:
    if not st.session_state.history:
        st.info("Zatím žádná uložená hodnocení. Vyhodnoťte odpověď a klikněte na **🔬 Vyhodnotit** v levém panelu.")
    else:
        st.markdown(f"**Celkem uložených hodnocení: {len(st.session_state.history)}**")

        # Table
        chips = {"CR":"chip-CR","CRi":"chip-CRi","PR":"chip-PR","PR-L":"chip-PRL",
                 "nPR":"chip-nPR","SD":"chip-SD","PD":"chip-PD"}
        rows=""
        for i,e in enumerate(st.session_state.history):
            chip_cls=chips.get(e["response"],"chip-SD")
            spd_s=f"{e['spd_pct']:+.1f}%" if e.get('spd_pct') is not None else "—"
            rows+=f"""<tr>
              <td>{i+1}</td>
              <td><b>{e['date']}</b></td>
              <td>{e['tp']}</td>
              <td><span class="{chip_cls}">{e['response']}</span></td>
              <td>{spd_s}</td>
              <td>{e['lymph_curr'] or '—'}</td>
              <td>{e['hgb_curr'] or '—'}</td>
              <td>{e['plt_curr'] or '—'}</td>
            </tr>"""

        st.markdown(f"""
        <table class="hist-table">
          <tr><th>#</th><th>Datum</th><th>Timepoint</th><th>Odpověď</th><th>ΔSPD</th><th>Lymfocyty</th><th>Hgb</th><th>PLT</th></tr>
          {rows}
        </table>
        """, unsafe_allow_html=True)

        # SPD trend chart
        try:
            import plotly.graph_objects as go
            spd_data = [(e["date"],e["tp"],e["spd_pct"]) for e in st.session_state.history if e.get("spd_pct") is not None]
            if spd_data:
                st.markdown("#### Trend SPD")
                labels=[f"{d[1]} ({d[0]})" for d in spd_data]
                vals=[d[2] for d in spd_data]
                colors=["#16a34a" if v<=-50 else ("#dc2626" if v>=50 else "#ca8a04") for v in vals]
                fig=go.Figure()
                fig.add_trace(go.Bar(x=labels,y=vals,marker_color=colors,text=[f"{v:+.1f}%" for v in vals],textposition="outside"))
                fig.add_hline(y=-50,line_dash="dash",line_color="#16a34a",annotation_text="≥50% pokles (PR threshold)")
                fig.add_hline(y=50, line_dash="dash",line_color="#dc2626",annotation_text="≥50% nárůst (PD threshold)")
                fig.update_layout(yaxis_title="Změna SPD (%)",xaxis_title="",plot_bgcolor="white",
                                  paper_bgcolor="white",height=350,margin=dict(t=20,b=20))
                st.plotly_chart(fig,use_container_width=True)

            # Response timeline
            resp_vals={"CR":6,"CRi":5,"PR":4,"PR-L":3,"nPR":4,"SD":2,"PD":1}
            resp_colors={"CR":"#16a34a","CRi":"#2563eb","PR":"#3b82f6","PR-L":"#9333ea","nPR":"#0284c7","SD":"#ca8a04","PD":"#dc2626"}
            r_labels=[f"{e['tp']} ({e['date']})" for e in st.session_state.history]
            r_vals=[resp_vals.get(e["response"],2) for e in st.session_state.history]
            r_colors=[resp_colors.get(e["response"],"#ca8a04") for e in st.session_state.history]
            r_text=[e["response"] for e in st.session_state.history]

            if len(st.session_state.history)>1:
                st.markdown("#### Vývoj odpovědi")
                fig2=go.Figure()
                fig2.add_trace(go.Scatter(x=r_labels,y=r_vals,mode="lines+markers+text",
                    marker=dict(color=r_colors,size=16),line=dict(color="#94a3b8",width=2),
                    text=r_text,textposition="top center",textfont=dict(size=12,color=r_colors)))
                fig2.update_layout(
                    yaxis=dict(tickvals=list(resp_vals.values()),ticktext=list(resp_vals.keys()),
                               title="",range=[0,7]),
                    xaxis_title="",plot_bgcolor="white",paper_bgcolor="white",height=300,
                    margin=dict(t=20,b=20),showlegend=False)
                st.plotly_chart(fig2,use_container_width=True)
        except ImportError:
            st.caption("Pro grafy nainstalujte: pip install plotly")

        if st.button("🗑 Vymazat historii"):
            st.session_state.history=[]
            st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
#  TAB 3 — Export
# ─────────────────────────────────────────────────────────────────────────────
with tab_export:
    st.markdown("### 📄 Export hodnocení")

    if not st.session_state.result:
        st.info("Nejprve proveďte hodnocení v záložce Aktuální hodnocení.")
    else:
        res2 = st.session_state.result
        export_type = st.radio("Formát exportu", ["DOCX (Word)", "JSON (data)"], horizontal=True)

        if export_type=="DOCX (Word)":
            try:
                from docx import Document as DocxDoc
                from docx.shared import Pt, RGBColor, Inches, Cm
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                import copy

                def set_cell_bg(cell, hex_color):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), hex_color)
                    tcPr.append(shd)

                def add_run(para, text, bold=False, size=11, color=None):
                    run=para.add_run(text)
                    run.bold=bold
                    run.font.size=Pt(size)
                    run.font.name="Calibri"
                    if color: run.font.color.rgb=RGBColor(*color)
                    return run

                doc=DocxDoc()

                # Margins
                for sec in doc.sections:
                    sec.top_margin=Cm(2)
                    sec.bottom_margin=Cm(2)
                    sec.left_margin=Cm(2.5)
                    sec.right_margin=Cm(2.5)

                # Header line
                h=doc.add_paragraph()
                add_run(h,"CLL RESPONSE ASSESSMENT",bold=True,size=16,color=(16,42,90))
                h.alignment=WD_ALIGN_PARAGRAPH.LEFT
                hpf=h._p.get_or_add_pPr()
                pb=OxmlElement('w:pBdr')
                bot=OxmlElement('w:bottom')
                bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
                bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'102A5A')
                pb.append(bot); hpf.append(pb)

                sub=doc.add_paragraph()
                add_run(sub,"Hodnocení odpovědi na léčbu dle iwCLL 2018 Guidelines",size=10,color=(100,100,100))

                doc.add_paragraph()

                # Patient info table
                info_table=doc.add_table(rows=2,cols=4)
                info_table.style='Table Grid'
                headers=["Pacient / ID","Studie / Protokol","Timepoint","Datum hodnocení"]
                values=[pat_id or "—", study_name or "—", tp_label or "—", eval_date.strftime("%-d. %-m. %Y")]
                for i,(h2,v) in enumerate(zip(headers,values)):
                    hc2=info_table.rows[0].cells[i]
                    vc=info_table.rows[1].cells[i]
                    set_cell_bg(hc2,"102A5A")
                    hp2=hc2.paragraphs[0]
                    add_run(hp2,h2,bold=True,size=9,color=(255,255,255))
                    vp=vc.paragraphs[0]
                    add_run(vp,v,size=10)

                doc.add_paragraph()

                # Result
                resp_colors_docx={
                    "CR":(21,128,61),"CRi":(30,58,138),"PR":(29,78,216),
                    "PR-L":(109,40,217),"nPR":(3,105,161),"SD":(133,77,14),"PD":(185,28,28)
                }
                rc=resp_colors_docx.get(res2["response"],(100,100,100))
                rp=doc.add_paragraph()
                add_run(rp,f"VÝSLEDEK: {res2['response']}",bold=True,size=22,color=rc)
                rp.alignment=WD_ALIGN_PARAGRAPH.CENTER

                dp=doc.add_paragraph()
                add_run(dp,res2["desc"],size=11)
                dp.alignment=WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()

                # Summary bar
                sp2=doc.add_paragraph()
                add_run(sp2,f"Patologické při baseline:  A: {res2['A_abn']}  |  B: {res2['B_abn']}",size=10,color=(80,80,80))
                sp3=doc.add_paragraph()
                add_run(sp3,f"Splněná zlepšení:  A: {res2['A_met']}/{res2['A_abn']}  |  B: {res2['B_met']}/{res2['B_abn']}",size=10,color=(80,80,80))

                doc.add_paragraph()

                # Parameters table
                doc.add_paragraph().add_run("DETAILNÍ ROZBOR PARAMETRŮ").bold=True
                all_items=[("A",n,v,s) for n,v,s,_ in res2["a_items"]] + \
                           [("B",n,v,s) for n,v,s,_ in res2["b_items"]]

                tbl=doc.add_table(rows=1,cols=4)
                tbl.style='Table Grid'
                hdr_row=tbl.rows[0]
                for i2,(hx) in enumerate(["Skupina","Parametr","Hodnota","Hodnocení"]):
                    set_cell_bg(hdr_row.cells[i2],"102A5A")
                    add_run(hdr_row.cells[i2].paragraphs[0],hx,bold=True,size=9,color=(255,255,255))

                status_labels={"met":"✓ Zlepšeno","notmet":"✗ Nezlepšeno","warn":"⚠ Pozor","na":"— N/A"}
                status_colors={"met":(21,128,61),"notmet":(185,28,28),"warn":(133,77,14),"na":(100,100,100)}
                for idx,(grp,nm,vl,st2) in enumerate(all_items):
                    row=tbl.add_row()
                    bg="F2F4F8" if idx%2==0 else "FFFFFF"
                    for ci in range(4): set_cell_bg(row.cells[ci],bg)
                    add_run(row.cells[0].paragraphs[0],grp,size=10)
                    add_run(row.cells[1].paragraphs[0],nm,size=10)
                    add_run(row.cells[2].paragraphs[0],str(vl),size=10)
                    sc=status_colors.get(st2,(100,100,100))
                    add_run(row.cells[3].paragraphs[0],status_labels.get(st2,"—"),size=10,color=sc)

                doc.add_paragraph()

                # Notes
                if res2["notes"]:
                    np2=doc.add_paragraph()
                    add_run(np2,"POZNÁMKY",bold=True,size=11,color=(192,57,43))
                    for n2 in res2["notes"]:
                        doc.add_paragraph(f"⚠ {n2}",style="List Bullet")

                doc.add_paragraph()

                # Doc text
                doc.add_paragraph().add_run("TEXT PRO DOKUMENTACI").bold=True
                dp2=doc.add_paragraph()
                add_run(dp2,res2["doc_text"],size=11)

                # Footer
                doc.add_paragraph()
                fp=doc.add_paragraph()
                add_run(fp,"Hodnocení provedeno dle iwCLL 2018 guidelines. Kalkulátor slouží jako pomůcka — klinické rozhodnutí závisí na lékaři.",
                        size=9,color=(150,150,150))

                buf=io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                fn=f"CLL_hodnoceni_{(pat_id or 'pacient').replace(' ','_')}_{eval_date}.docx"
                st.download_button("⬇️ Stáhnout DOCX",data=buf,file_name=fn,
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.success("Dokument připraven ke stažení.")
            except ImportError:
                st.error("Nainstalujte: pip install python-docx")
            except Exception as e:
                st.error(f"Chyba při generování DOCX: {e}")

        else:
            # JSON export
            export_data={
                "patient": pat_id, "study": study_name, "timepoint": tp_label,
                "date": str(eval_date), "response": res2["response"],
                "description": res2["desc"], "A_met": res2["A_met"], "A_abn": res2["A_abn"],
                "B_met": res2["B_met"], "B_abn": res2["B_abn"], "notes": res2["notes"],
                "doc_text": res2["doc_text"], "history": st.session_state.history
            }
            json_str=json.dumps(export_data,ensure_ascii=False,indent=2)
            st.download_button("⬇️ Stáhnout JSON",data=json_str,
                               file_name=f"CLL_{pat_id or 'export'}_{eval_date}.json",mime="application/json")
            st.code(json_str,language="json")
